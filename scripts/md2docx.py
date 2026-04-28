import sys
import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def convert_md_to_docx(md_path: str, docx_path: str):
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    doc = Document()

    lines = content.split('\n')
    in_code_block = False

    for line in lines:
        line = line.rstrip()

        if line.startswith('```'):
            in_code_block = not in_code_block
            continue

        if in_code_block:
            para = doc.add_paragraph(line, style='Normal')
            continue

        if line.startswith('# '):
            para = doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            para = doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            para = doc.add_heading(line[4:], level=3)
        elif line.startswith('- '):
            para = doc.add_paragraph(line[2:], style='List Bullet')
        elif line.startswith('| '):
            if '---' in line:
                continue
            parts = [p.strip() for p in line.split('|')[1:-1]]
            para = doc.add_paragraph(' | '.join(parts))
        elif line == '---':
            continue
        elif line.strip():
            para = doc.add_paragraph(line)

    doc.save(docx_path)
    print(f"Created: {docx_path}")

if __name__ == '__main__':
    convert_md_to_docx('docs/BRD-skills-system.md', 'docs/BRD-skills-system.docx')
    convert_md_to_docx('docs/SDD-skills-system.md', 'docs/SDD-skills-system.docx')